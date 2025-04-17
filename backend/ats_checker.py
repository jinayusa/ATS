from docx import Document
import io

ATS_SECTIONS = ["education", "experience", "skills", "projects", "certifications"]

def analyze_ats_format(file_bytes, filename):
    issues = []
    sections_found = set()

    if filename.endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))

        # Check for tables
        if doc.tables:
            issues.append("Tables detected — ATS may skip content in tables.")

        # Check for images
        if any(shape for shape in doc.inline_shapes):
            issues.append("Images detected — ATS cannot read images.")

        # Check for standard sections
        for para in doc.paragraphs:
            text = para.text.strip().lower()
            for section in ATS_SECTIONS:
                if section in text:
                    sections_found.add(section)

        missing_sections = set(ATS_SECTIONS) - sections_found
        if missing_sections:
            issues.append(f"Missing recommended sections: {', '.join(missing_sections)}.")

    else:
        issues.append("ATS formatting check only available for .docx resumes.")

    ats_score = 100
    if issues:
        ats_score = max(0, 100 - len(issues) * 20)

    return {
        "ats_score": ats_score,
        "format_issues": issues
    }
